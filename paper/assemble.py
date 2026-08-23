"""Assemble paper sections (markdown subset, see STYLE.md) into an ICLR-like .docx with figures, tables, rendered equations."""
import re, os, sys, glob, io
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
plt.rcParams["mathtext.fontset"] = "stix"; plt.rcParams["font.family"] = "STIXGeneral"

P = os.path.dirname(os.path.abspath(__file__))
SEC = sorted(glob.glob(os.environ.get("SEC_DIR", f"{P}/sections") + "/S*.md"))
OUT = os.environ.get("OUT", f"{P}/paper.docx")
TITLE = "Few-Step Distillation of Multi-View RGB-D World Models: Does Geometry Survive?"
AUTHOR = "Tae Yang Hong"
AFFIL = "NAIS Lab"
EMAIL = "hongtaeyang1231@gmail.com"

doc = Document()
for s in doc.sections:
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.left_margin = s.right_margin = Inches(1.0); s.top_margin = s.bottom_margin = Inches(1.0)
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.08
for lvl, size in ((1, 13), (2, 11.5), (3, 10.5)):
    h = doc.styles[f"Heading {lvl}"]; h.font.name = "Times New Roman"; h.font.size = Pt(size); h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0); h.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h.paragraph_format.space_before = Pt(12 if lvl == 1 else 8); h.paragraph_format.space_after = Pt(4)

SYM = {r"\rho": "ρ", r"\Omega": "Ω", r"\omega": "ω", r"\gamma": "γ", r"\Gamma": "Γ", r"\kappa": "κ", r"\tau": "τ", r"\pi": "π", r"\xi": "ξ", r"\eta": "η", r"\zeta": "ζ", r"\psi": "ψ", r"\Psi": "Ψ", r"\Sigma": "Σ", r"\Theta": "Θ", r"\Lambda": "Λ", r"\Phi": "Φ", r"\nu": "ν", r"\chi": "χ", r"\iota": "ι", r"\upsilon": "υ", r"\Pi": "Π", r"\ell": "ℓ", r"\partial": "∂", r"\bar{z}": "z̄", r"\hat{x}": "x̂", r"\tilde{z}": "z̃", r"\sigma": "σ", r"\epsilon": "ε", r"\varepsilon": "ε", r"\theta": "θ", r"\phi": "φ", r"\lambda": "λ", r"\beta": "β", r"\alpha": "α",
       r"\mu": "μ", r"\delta": "δ", r"\Delta": "Δ", r"\nabla": "∇", r"\times": "×", r"\cdot": "·", r"\approx": "≈", r"\leq": "≤", r"\geq": "≥",
       r"\rightarrow": "→", r"\to": "→", r"\in": "∈", r"\sim": "~", r"\infty": "∞", r"\pm": "±", r"\mathbb{E}": "𝔼", r"\mathbb{R}": "ℝ",
       r"\ldots": "…", r"\dots": "…", r"\|": "‖", r"\langle": "⟨", r"\rangle": "⟩", r"\log": "log", r"\exp": "exp", r"\min": "min", r"\max": "max",
       r"\operatorname{median}": "median", r"\mathrm{median}": "median", r"\text{median}": "median", r"\circ": "∘", r"\star": "★", r"\ast": "∗",
       r"\hat": "^", r"\bar": "‾", r"\tilde": "~", r"\prime": "′", r"\partial": "∂", r"\sum": "Σ", r"\int": "∫", r"\sqrt": "√", r"\propto": "∝", r"\ne": "≠", r"\neq": "≠"}
SUB = str.maketrans("0123456789+-=()aeijklmnoprstuvx", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ")
SUP = str.maketrans("0123456789+-=()n", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ")

def tex_to_text(t):
    t = t.strip()
    for k, v in sorted(SYM.items(), key=lambda kv: -len(kv[0])):
        t = t.replace(k, v)
    t = re.sub(r"\\(mathbf|mathrm|text|boldsymbol|mathcal|textrm|operatorname)\{([^{}]*)\}", r"\2", t)
    t = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"(\1)/(\2)", t)
    def sub(m):
        s = m.group(1); return s.translate(SUB) if all(c in "0123456789+-=()aeijklmnoprstuvx" for c in s) else "_" + s
    def sup(m):
        s = m.group(1); return s.translate(SUP) if all(c in "0123456789+-=()n" for c in s) else "^" + s
    t = re.sub(r"_\{([^{}]*)\}", sub, t); t = re.sub(r"_(\w)", sub, t)
    t = re.sub(r"\^\{([^{}]*)\}", sup, t); t = re.sub(r"\^(\w)", sup, t)
    t = re.sub(r"\\left|\\right|\\,|\\;|\\!|\\ ", "", t)
    t = t.replace("{", "").replace("}", "").replace("\\", "")
    return t

def render_eq(tex, tag, idx):
    """render display equation with mathtext; fallback to text"""
    path = f"{P}/figures/eq_{idx}.png"
    try:
        fig = plt.figure(figsize=(0.1, 0.1))
        fig.text(0, 0, f"${tex}$", fontsize=13, family="serif")
        fig.savefig(path, dpi=300, bbox_inches="tight", pad_inches=0.04, transparent=False, facecolor="white")
        plt.close(fig)
        return path
    except Exception as e:
        print(f"  [eq {idx}] mathtext failed: {str(e)[:80]} -> text fallback"); plt.close("all")
        return None

def add_runs(par, text):
    """inline: **bold**, *italic*, `code`, $math$"""
    pos = 0
    for m in re.finditer(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$[^$]+\$)", text):
        if m.start() > pos: par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"): r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"): r = par.add_run(tok[1:-1]); r.font.name = "Courier New"; r.font.size = Pt(9.5)
        elif tok.startswith("$"): r = par.add_run(tex_to_text(tok[1:-1])); r.italic = True
        else: r = par.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text): par.add_run(text[pos:])

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr(); borders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom"):
        if edge in kwargs:
            el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(kwargs[edge])); el.set(qn("w:color"), "000000"); borders.append(el)
    tcPr.append(borders)

def add_table(rows, caption):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = re.match(r"(Table \d+:)\s*(.*)", caption)
    if m:
        r = cap.add_run(m.group(1) + " "); r.bold = True; r.font.size = Pt(9.5); add_runs(cap, m.group(2))
        for r in cap.runs: r.font.size = Pt(9.5)
    else:
        add_runs(cap, caption)
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    # remove grid borders, keep booktabs-like rules
    tbl = t._tbl; tblPr = tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "nil"); b.append(el)
    tblPr.append(b)
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.cell(i, j); cell.text = ""
            par = cell.paragraphs[0]; par.paragraph_format.space_after = Pt(0); par.paragraph_format.space_before = Pt(0)
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            txt = row[j] if j < len(row) else ""
            add_runs(par, txt)
            for r in par.runs:
                r.font.size = Pt(8.5)
                if i == 0: r.bold = True
            if i == 0: set_cell_border(cell, top=8, bottom=4)
            if i == len(rows) - 1: set_cell_border(cell, bottom=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_figure(path, caption):
    full = os.path.join(P, path)
    if not os.path.exists(full):
        p = doc.add_paragraph(f"[missing figure: {path}]"); p.runs[0].font.color.rgb = RGBColor(200, 0, 0); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=Inches(6.3))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    m = re.match(r"(Figure \d+:)\s*(.*)", caption)
    if m:
        r = cap.add_run(m.group(1) + " "); r.bold = True
        add_runs(cap, m.group(2))
    else:
        add_runs(cap, caption)
    for r in cap.runs: r.font.size = Pt(9.5)
    cap.paragraph_format.space_after = Pt(10)

# ── title block
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run(TITLE); r.bold = True; r.font.size = Pt(16)
ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ap.add_run(AUTHOR); r.bold = True; r.font.size = Pt(11)
ap.add_run(f"\n{AFFIL}\n{EMAIL}").font.size = Pt(10)
doc.add_paragraph()

eq_idx = 0
for f in SEC:
    lines = open(f, encoding="utf-8").read().split("\n")
    i = 0; pending_caption = None
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        if ln.startswith("# "):
            h = ln[2:].strip()
            if h.lower() == "abstract":
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Abstract"); r.bold = True; r.font.size = Pt(12)
            else:
                doc.add_heading(h, level=1)
            i += 1; continue
        if ln.startswith("## "):
            doc.add_heading(ln[3:].strip(), level=2); i += 1; continue
        if ln.startswith("### "):
            doc.add_heading(ln[4:].strip(), level=3); i += 1; continue
        m = re.match(r"!\[(.*?)\]\((.*?)\)", ln.strip())
        if m:
            add_figure(m.group(2), m.group(1)); i += 1; continue
        if re.match(r"^Table \d+:", ln.strip()) and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            pending_caption = ln.strip(); i += 1; continue
        if ln.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            add_table(rows, pending_caption or ""); pending_caption = None; continue
        if ln.strip().startswith("$$"):
            buf = ln.strip()
            while not (buf.count("$$") >= 2 and buf.rstrip().endswith(("$$",)) or re.search(r"\$\$\s*\(\d+\)\s*$", buf)) and i + 1 < len(lines):
                i += 1; buf += " " + lines[i].strip()
            mm = re.match(r"\$\$(.*)\$\$\s*(\(\d+\))?\s*$", buf, re.S)
            tex, tag = (mm.group(1).strip(), mm.group(2) or "") if mm else (buf.strip("$ "), "")
            eq_idx += 1
            img = render_eq(tex, tag, eq_idx)
            t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
            c0, c1 = t.cell(0, 0), t.cell(0, 1); c0.width = Inches(5.6); c1.width = Inches(0.7)
            p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img:
                from PIL import Image as _I
                w_, h_ = _I.open(img).size; hin = h_ / 300 * 1.15; win = w_ / 300 * 1.15
                if win > 5.4: hin *= 5.4 / win
                p0.add_run().add_picture(img, height=Inches(hin))
            else: p0.add_run(tex_to_text(tex)).italic = True
            p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p1.add_run(tag)
            i += 1; continue
        if ln.strip().startswith("```"):
            i += 1; code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph(); r = p.add_run("\n".join(code)); r.font.name = "Courier New"; r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Inches(0.3); continue
        if ln.strip().startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet"); add_runs(p, lines[i].strip()[2:]); i += 1
            continue
        # paragraph: gather until blank line
        buf = [ln.strip()]
        while i + 1 < len(lines) and lines[i + 1].strip() and not re.match(r"^(#|\||!\[|\$\$|```|- |Table \d+:)", lines[i + 1].strip()):
            i += 1; buf.append(lines[i].strip())
        text = " ".join(buf)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if os.path.basename(f).startswith("S6"):
            p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.first_line_indent = Inches(-0.3); p.paragraph_format.space_after = Pt(3)
            for r_ in p.runs: r_.font.size = Pt(9.5)
        add_runs(p, text)
        if os.path.basename(f).startswith("S6"):
            for r_ in p.runs: r_.font.size = Pt(9.5)
        i += 1

doc.save(OUT)
print("saved", OUT, "equations", eq_idx)
